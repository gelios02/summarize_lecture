import os
import re
from datetime import datetime
from time import sleep
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from bs4 import BeautifulSoup
import markdown
import win32com.client


# Инициализация GigaChat
chat = GigaChat(
    credentials='NWMwMmQxMTQtMjM0Mi00MGQzLTk0ZWYtOTk1NDY2ODYyZmJiOjViZmQ0MGI2LTY1NTEtNDY1Mi05YTVkLWFkYTQ0MWQzYTlhMQ==',
    verify_ssl_certs=False,
    model='GigaChat-Max'
)

MAX_TEXT_LENGTH = 3000  # Максимальная длина текста для одного запроса
RETRY_LIMIT = 10  # Количество повторных попыток
DOCX_SAVE_PATH = os.path.abspath("D://lemon/summarize_lecture/saved_lecture")  # Папка для сохранения DOCX


def clean_directory(path):
    """
    Очистка указанной папки.
    """
    if os.path.exists(path):
        for file in os.listdir(path):
            file_path = os.path.join(path, file)
            os.remove(file_path)
    else:
        os.makedirs(path)


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Конвертация DOCX в PDF с помощью pywin32.
    """
    import pythoncom
    pythoncom.CoInitialize()  # Инициализация COM
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 — формат PDF
        doc.Close()
        word.Quit()
    finally:
        pythoncom.CoUninitialize()  # Завершаем работу с COM



def save_to_docx(content, path):
    """
    Сохранение текста с Markdown-разметкой в формате DOCX (Microsoft Word).
    """
    # Создаём документ Word
    doc = Document()
    # Устанавливаем шрифт и размер текста по умолчанию
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    doc.styles['Normal'].font.size = Pt(12)

    # Добавляем заголовок документа
    title = doc.add_heading("Лекция", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 139)  # Синий цвет заголовка

    # Преобразуем Markdown в HTML
    html = markdown.markdown(content)
    soup = BeautifulSoup(html, 'html.parser')

    # Сохраняем список уже добавленных заголовков, чтобы исключить дублирование
    added_headings = set()

    # Обрабатываем HTML и добавляем элементы в Word
    for element in soup.descendants:
        if element.name in ['h1', 'h2', 'h3']:
            heading_text = element.get_text().strip()
            if heading_text not in added_headings:
                level = {'h1': 1, 'h2': 2, 'h3': 3}[element.name]
                heading = doc.add_heading(heading_text, level=level)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                added_headings.add(heading_text)

                # Устанавливаем отступы для заголовков
                paragraph_format = heading.paragraph_format
                paragraph_format.left_indent = Pt(18)
                paragraph_format.hanging_indent = Pt(36)

        # Стили для сильного выделения
        elif element.name == 'strong':
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(18)
            paragraph_format.hanging_indent = Pt(36)
            run = paragraph.add_run(element.get_text())
            run.bold = True

        # Стили для курсивного текста
        elif element.name == 'em':
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(18)
            paragraph_format.hanging_indent = Pt(36)
            run = paragraph.add_run(element.get_text())
            run.italic = True

        # Параграфы с отступами
        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(18)
            paragraph_format.hanging_indent = Pt(36)

        # Маркированные списки
        elif element.name == 'ul':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(style='List Bullet')
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Pt(18)
                paragraph_format.hanging_indent = Pt(36)
                paragraph.add_run(li.get_text())

        # Нумерованные списки
        elif element.name == 'ol':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(style='List Number')
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Pt(18)
                paragraph_format.hanging_indent = Pt(36)
                paragraph.add_run(li.get_text())

        # Горизонтальная линия
        elif element.name == 'hr':
            paragraph = doc.add_paragraph()
            paragraph.add_run("—" * 30)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(18)
            paragraph_format.hanging_indent = Pt(36)

    # Сохраняем документ
    docx_file = os.path.join(path, "summarized_lecture_docx.docx")
    doc.save(docx_file)
    print(f"DOCX сохранён по пути: {docx_file}")
    return docx_file





def retry_function(func, retries=RETRY_LIMIT, *args, **kwargs):
    for attempt in range(retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            print(f"Ошибка в {func.__name__}: {e}. Попытка {attempt + 1} из {retries}...")
            sleep(2)
    return None


def validate_detail_level(detail_level):
    return isinstance(detail_level, int) and 1 <= detail_level <= 10


def split_text(text, max_length):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    parts = []
    current_part = ""

    for sentence in sentences:
        if len(current_part) + len(sentence) <= max_length:
            current_part += sentence + " "
        else:
            parts.append(current_part.strip())
            current_part = sentence + " "

    if current_part:
        parts.append(current_part.strip())

    return parts


def summarize_text(text, detail_level):
    if not text.strip():
        return "Ошибка: текст для суммаризации пустой."

    if not validate_detail_level(detail_level):
        return "Ошибка: уровень подробности должен быть числом от 1 до 10."

    text_parts = split_text(text, MAX_TEXT_LENGTH)
    summarized_parts = []

    for part in text_parts:
        prompt = (
            "Проанализируй текст лекции и выпиши только основные и важные моменты. "
            "Избегай излишних деталей и формулировок, которые не относятся к сути лекции.  Убери любые приветствия, прощания или метаданные."
            "Структурируй текст так, чтобы он оставался понятным. Если текст на английском, ответ должен быть на английском, напиши мне мою структурированную лекцию на английском языке. Если текст на русском, ответ должен быть на русском."
            "Для английского текста добавь минимальные пояснения на русском для повышения понимания."
            "Если подробность = 1, напиши текст максимально кратко, только самые главные тезисы.\n"
            "Если подробность = 10, очень подробно изложи содержание, но без лишней воды.\n"
            f"Уровень подробности: {detail_level}\n"
            "Вот текст лекции: " + part
        )

        system_message = SystemMessage(
            content="Ты ассистент, который умеет анализировать и суммаризировать текстовые лекции. "
                    "Если лекция на английском, пиши её на английском, добавляя минимальные пояснения на русском."
        )
        human_message = HumanMessage(content=prompt)

        def request_summary():
            response = chat.invoke([system_message, human_message])
            return response.content.strip()

        summary = retry_function(request_summary)

        if not summary:
            return "Ошибка: не удалось получить суммаризацию текста после нескольких попыток."

        summarized_parts.append(summary)

    structured_summary = "\n\n".join(summarized_parts)
    return structured_summary.strip()


def generate_additional_materials(summary):
    prompt = (
        "На основе следующего суммаризированного текста лекции добавь обучающие материалы и пояснения для лучшего понимания. "
        "Если текст на английском, используй английский язык для примеров и пояснений, добавив минимум комментариев на русском. Для русского текста придерживайся формата подробных пояснений на русском. Материалы должны быть структурированы и понятны.\n"
        "Вот суммаризированный текст лекции: " + summary
    )

    system_message = SystemMessage(
        content="Ты помощник, который создает обучающие материалы на основе предоставленных текстов."
    )
    human_message = HumanMessage(content=prompt)

    def request_additional_materials():
        response = chat.invoke([system_message, human_message])
        return response.content.strip()

    additional_materials = retry_function(request_additional_materials)

    if not additional_materials:
        return "Ошибка: не удалось сгенерировать дополнительные материалы."

    return additional_materials

def replace_numbered_lists_with_alternative_format(text):

    # Паттерн для поиска числовых списков
    numbered_list_pattern = r"^\d+\.\s"

    # Альтернативный формат для списков (например, '- ' или буквенный формат 'a.', 'b.')
    def replacement(match):
        return "- "

    # Применяем замену
    processed_text = re.sub(numbered_list_pattern, replacement, text, flags=re.MULTILINE)
    return processed_text
def process_lecture(text, detail_level, smart_generation=False):
    # Очищаем папку
    clean_directory(DOCX_SAVE_PATH)
    summary = summarize_text(text, detail_level)

    if "Ошибка" in summary:
        return summary

    if smart_generation:
        additional_materials = generate_additional_materials(summary)
        if "Ошибка" not in additional_materials:
            final_text = f"{summary}\n\n{additional_materials}"
        else:
            final_text = summary
    else:
        final_text = summary

    final_text = replace_numbered_lists_with_alternative_format(final_text)
    docx_file = save_to_docx(final_text, DOCX_SAVE_PATH)
    pdf_file = os.path.join(DOCX_SAVE_PATH, "summarized_lecture_pdf.pdf")
    convert_docx_to_pdf(docx_file, pdf_file)
    print(f"PDF сохранён по пути: {pdf_file}")

    return final_text



# if __name__ == "__main__":
#     lecture_text = """Hello, I welcome to this Learn English Elementary Recording bruute to you by the British Council. To find that more and the access language activaties and audioscripts visit or learn English Website от ww. Britishcansole.org foredlearn English. Hello, and welcome to Learn English Elementary Podcast Number 1. My names Rawy. And I'm Tess. We are presenters and we go lots of things be you to listing today. The before we starts, I think we should interduce authes. Rawy? И я is how alway? No have your business Rabbie. And she love sdanceling and writing her mountain bike, ok? Okay. And this is Rabbie, he comes from Manchester, he s 23. Oh, are you? Oh, yes. He's a great cook. Thanks? And there's warm more person for eature it am note like 20 jews a perjewser golden. Show today, и let's start with a fast action code, I like to made. We are a sing a show few today, and it start with a fast action could I like to made. We are a people a simple question which famous person that or alive would you like to made? И я Анджелина Джоли. И я и я It's usually just music videos and things, back is famous and beautiful, and people want to see her, so she can get a lot of attention for the things, что she want to change. Do you like a films? Time. And you give is miniantes as you can in 10 secondies. Art 2 place today are dannio? Dannio 6thin and comes from London. Helloy, Dannio! Hi! And Alice! Alice is alsoy 6thin and she comes from Livpeople! We have you a topic and you have to right down all the word, you can think of. For example: We say bathroom and you write a list. Bath sope. Шампунь we'n't so 1. I'es quick it you can. Got the idea? Yeah, good. Okay, let's play. You could 10 seconds to right down. О и я life in the country their in. Today will here from Mike Sowden. MIKE is our man in New York. Нью-Йорк. What do you think of skyscrapers, taxis, noise, people? Джон Леннон сед, that new york is the capital city of the world, and it surt ney fills that way, but in the center of the exciteing noisy pluted city there is a place to find piece and quired. A short work from busy 50th Avenue in Manhattan you find Central Park. Central Park is the green heart of New York City. Ad over 25 млн. People come here HR to escape the city, people run, swim, clim, clim, or simply sit and readabock in the hugee parks differenceinary. And it's not just for people. Lut's of rare birds of made their home in the Park And their are legends of big cats heiding in the dark trees. All the wed locks completely natural, the Park is man made. Created over hundred 50 years ego. Фредерик Ло Onstead and Calvard Vaw дизайнed the Park это time, end the City was a very crowded, dirty and un heathy place to live. They wanted to make a place, were rich and pore people could find fresh air in the Dirty City. It does the same job today, as the friend of my always telse me: life in New York with be impossible without Central Park. Вау! Рили инетерестинг! We have from another wannather wannather wanna have a people next time. О, если я хочу дать о райтен и тэла something interesting about yoursity to your tell, we have to have from you. Новых An te the day were going to ask a question about celebrities. Who permote charities an ask people take it the money or raid. But the beginning of the show Willessen и Zara from Bristall Talking about Анджелины Джоли. Zara talks by Angelena Jollies work for Charity, and now a days lots a famous people, celebrities do the same. But is it a good idea? So teday's question is? И и я и я и я и я It's a big adventure for her. She is going to live, study and she hopes, have a going time here in UK. And we are going with the Carolin is going to studying New Castle in the North East of England. She speke s very good English, that this have first visit to Britain, so something is a very strienged for her We goting to follow Carolina Reanal Podcasts, anlysten to conversations, that he has in lots of different situations. И я и я и я Chicken works into a library. Chicken works into library. Works upter the counter and saest to the librarian. The librarian games to chicken to books. Book book. Hm, this is strange, so she disize to follow the chicken. She goes after the library and follows the chicken. The chicken cross is a road, walks on the street, turns the conna untily comes to the lake. Readit. And that the end of this parts of the show We have to go now but don't go away. After the this little break you going to hear Thom are English teacher. After every Том толкс оберт the language you heard and gives you ideas to help you learn. So don't go away, but will say gaby now. See you next time. Bye. Don't ticket to send us your emails. Here's a adress 1 more time. It's learning English poccause от britishcauncil.org. You are listening to a learning English elementary recoarding from the Britishcauncil. Hi, I'm Thom. And the end of every podcast you here from me. I'm going to talk about about some of the language you heard in the programs, and talk about ways to help you learning English. Remember Carolina in the Appold. Listened to part of her conversation, again. Is this your first visit to the UK? Yes, it is. And what to you going to do here? I'm a student, I'm going to study at the University of Newcastle. Carolina and the Immigration ofice a talking about the future. Carolina's time in Britain. To talk about the future, they book use going to. The Immigration offers is is and what are you going to do here? And Carolina's is I'm going to study at the University of New Council. They boke used going to to talk about the future, because they are talking about plans. When you immigration officer sase: What are you going to do here? He's asking Carolina what her plan is. And she saes I'm going to study, because that's her plan. She diside it before she left Venes Wheeler. So, will use going to to talk about future plans. But listing twoy noter put of Carolina's conversation. The Immigration offers a said: How long do you intent to stay in the country? And Do you intent to working this country? Again, He was asking about Carolina's future plans. Body said Do you intent? Intent is a forma way to talk or ask about plans. You might here this verb intent at an airport immigration desk or on animmigration form. It's not a wait to ask about your plans. 1 more thing. Do you knowtased when Carolina arrived at the desk The Immigration of a suset good eaving? He didn't say good night. Do you know why not? We are ney say good night? We're we're we're good bye or when we're go to bead? We're we're we're go to bead? Okay. Ok. In another part of the show, We heard Daniel and Alice blanger agame. Listened to part of it again. You gott 10 seconds to right down Things you can find in a kitchen. Think about, how you put new words in the your notebook? Do you put theme in althabetically? All a words beginning with a, then all a words bigining with b? Do you organized your new words an all away? Some people put words into the notbooks in word familys, they put words together, that are connected in someway. You could have a page in over cabilary notebook cooled kitchen, and you could key all the words from the game, fridge, cooker, pan, all the those words on the kitchen page of yould outbook you could have pages for say sports, football, tennis, боулинг и some. And you can right more than just the words, you can right the verbs that go with the words, play, football, but go bowling, or go skeing. There's no right way or runway the keep you or new vercabulary. You have to find a way that helpse you remember the new words. Ok. Oh 1 but the way, is there are any words of the game, that you don't know, remember, that you can find am in the website. You can read all of the podcast, and if you click on the word, it it al take you to were dictionary, that tales you what the word means. Will give it at adress again at the end. So go on find a pen to There are 2 things, there I want to talk ab at. Firstly The old lady chouled Carolina Dear. She said: I don't know dear. Sometime is old the people my choul you dear. It's a friendly afaction think to do. But be caful it my sound a bit strange, if you tried to use it yearself. The other they I knotist was at the old lady said it depends? She did not answer, beca she needed more information. Can you translate it depends in the your language? Thaty use it in English this we'k. I'm got is top the I'll talk you all again next time. Remember, you can sand your questions to me от Learn English podcast от britishcancele.org I'll be happy to ones your questions. In a moment you al heed years for the website, we can read everything your heard in this podgast. Right, that's al for the this time. Bye fanow. See your next time. This recording was brute to you about at british canceule.
# """
#     detail = 10
#     smart = True
#
#     result = process_lecture(lecture_text, detail, smart)
#     print(result)
